"""
Module for generating PDF and DOCX versions of legal analysis results.
"""

import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt

# Font registration
FONT_PATH = "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf"
FONT_NAME = "NotoSans"

try:
    if os.path.exists(FONT_PATH):
        pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
    else:
        print(f"[WARNING] Devanagari font not found at {FONT_PATH}. Falling back to standard fonts.")
        FONT_NAME = "Helvetica"
except Exception as e:
    print(f"[ERROR] Font registration failed: {e}")
    FONT_NAME = "Helvetica"

def generate_pdf(data: dict) -> io.BytesIO:
    """Generate a professional PDF report from analysis results."""
    lang = "hi" if data.get("output_language") == "Hindi" else "en"
    # Use Unicode font for Hindi, standard or Unicode for English
    active_font = FONT_NAME if lang == "hi" or FONT_NAME != "Helvetica" else "Helvetica"
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    # Custom styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName=active_font, # Use registered font
        fontSize=22,
        spaceAfter=12,
        textColor=colors.HexColor("#6366f1")
    )
    section_style = ParagraphStyle(
        'SectionStyle',
        parent=styles['Heading2'],
        fontName=active_font, # Use registered font
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor("#1e293b")
    )
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName=active_font, # Use registered font
        fontSize=10,
        leading=14
    )
    
    elements = []
    
    # 1. Header
    elements.append(Paragraph("LexAnalytica Analysis Report", title_style))
    elements.append(Paragraph(f"Document: {data.get('filename', 'Unknown')}", body_style))
    elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    elements.append(Paragraph(f"Output Language: {data.get('output_language', 'English')}", body_style))
    elements.append(Spacer(1, 24))
    
    # 2. Executive Summary
    summary = data.get("summary", {})
    if summary.get("extractive") or summary.get("abstractive"):
        elements.append(Paragraph("Executive Summary", section_style))
        if summary.get("abstractive"):
            elements.append(Paragraph("<b>AI Synthesis:</b>", body_style))
            elements.append(Paragraph(summary["abstractive"], body_style))
            elements.append(Spacer(1, 12))
        if summary.get("extractive"):
            elements.append(Paragraph("<b>Key Findings (Extractive):</b>", body_style))
            elements.append(Paragraph(summary["extractive"], body_style))
        elements.append(Spacer(1, 12))
    
    # 3. Entities
    entities = data.get("entities", {})
    if any(entities.values()):
        elements.append(Paragraph("Identified Legal Entities", section_style))
        
        entity_data = [["Type", "Extraction"]]
        for key, label in [
            ("persons", "Persons"),
            ("organizations", "Organizations"),
            ("dates", "Dates"),
            ("locations", "Locations"),
            ("case_numbers", "Case Numbers"),
            ("law_sections", "Law Sections")
        ]:
            items = entities.get(key, [])
            if items:
                entity_data.append([label, ", ".join(items)])
        
        if len(entity_data) > 1:
            t = Table(entity_data, colWidths=[100, 350])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#475569")),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), active_font), # Use active_font for whole table
                ('FONTNAME', (0, 0), (-1, 0), f"{active_font}-Bold" if active_font == "Helvetica" else active_font),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0"))
            ]))
            elements.append(t)
        elements.append(Spacer(1, 12))
        
    # 4. Metrics
    elements.append(Paragraph("Document Metrics", section_style))
    metrics = [
        ["Metric", "Value"],
        ["Word Count", str(data.get("word_count", 0))],
        ["Sentence Count", str(data.get("sentence_count", 0))],
        ["Character Length", str(data.get("text_length", 0))]
    ]
    t_metrics = Table(metrics, colWidths=[150, 100])
    t_metrics.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
        ('FONTNAME', (0, 0), (-1, -1), active_font),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
    ]))
    elements.append(t_metrics)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx(data: dict) -> io.BytesIO:
    """Generate a formatted Word document from analysis results."""
    doc = Document()
    
    # Title
    title = doc.add_heading('LexAnalytica Analysis Report', 0)
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run(f"Document: {data.get('filename', 'Unknown')}\n").bold = True
    p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p.add_run(f"Output Language: {data.get('output_language', 'English')}")
    
    # Executive Summary
    summary = data.get("summary", {})
    if summary.get("extractive") or summary.get("abstractive"):
        doc.add_heading('Executive Summary', level=1)
        if summary.get("abstractive"):
            doc.add_heading('AI Synthesis', level=2)
            doc.add_paragraph(summary["abstractive"])
        if summary.get("extractive"):
            doc.add_heading('Key Findings (Extractive)', level=2)
            doc.add_paragraph(summary["extractive"])
            
    # Entities
    entities = data.get("entities", {})
    if any(entities.values()):
        doc.add_heading('Identified Legal Entities', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light List Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type'
        hdr_cells[1].text = 'Extraction'
        
        for key, label in [
            ("persons", "Persons"),
            ("organizations", "Organizations"),
            ("dates", "Dates"),
            ("locations", "Locations"),
            ("case_numbers", "Case Numbers"),
            ("law_sections", "Law Sections")
        ]:
            items = entities.get(key, [])
            if items:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = ", ".join(items)
                
    # Metrics
    doc.add_heading('Document Metrics', level=1)
    doc.add_paragraph(f"Word Count: {data.get('word_count', 0)}")
    doc.add_paragraph(f"Sentence Count: {data.get('sentence_count', 0)}")
    doc.add_paragraph(f"Character Length: {data.get('text_length', 0)}")
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "Generated by LexAnalytica Institutional Intelligence"
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
